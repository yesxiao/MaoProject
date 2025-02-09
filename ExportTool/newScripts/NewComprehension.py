import docx
from docx.shared import Pt
from docx.text.run import Run

import MyParagraph
from ContentTools import parse_doc

if __name__ == '__main__':
    doc_path = r"D:\猫猫\0905\肖鲲测试\【改5】【答案】【选择题】中国地理：河西走廊（1）1~100（题完）_河西走廊（1）1~100（题完）_理解题.docx"
    save_doc_path = r"D:\猫猫\0905\肖鲲测试\【改5】【答案】【选择题】中国地理：河西走廊（1）1~100（题完）_河西走廊（1）1~100（题完）_理解题[save].docx"
    doc_analyse = MyParagraph.DocAnalyse(doc_path)
    category = ["【答案】", "【知识点】", "【解析】", "【详解】", "【点睛】"]
    main:str = "main"
    que:str = "qus"
    option:str = "option"
    result_of_parse_doc = parse_doc(doc_analyse, ["【答案】", "【知识点】", "【解析】", "【详解】", "【点睛】"])
    items = []
    for d in result_of_parse_doc:
        item = {}
        items.append(item)
        item[main] = d[main]
        item[que] = []
        for k in d :
            if result_of_parse_doc.__contains__(k):
                item[k] = d[k]
            else:
                if k.startswith(que):
                    option_key:str = option + k.replace(que,"")
                    o = {k: d[k], option_key: d[option_key]}
                    item[que].append(o)
    doc = docx.api.Document()
    from docx.oxml.ns import qn
    for item in items:
        #写入主题
        main_list = item[main]
        pa = doc.add_paragraph()
        for p in main_list:
            for c in p.contents:
                run = pa.add_run()
                run.font.name = 'Times New Roman'
                # run_2.font.name = '楷体'  # 注：如果想要设置中文字体，需在前面加上这一句
                run.font.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
                # 设置字体大小
                run.font.size = Pt(14)
                tmp_content: str = c.content
                run.text = tmp_content
                if run.text.startswith("[path:\""):
                    path = tmp_content.replace("[path:\"", "").replace("\"]", "")
                    run.text = ""
                    r: Run = run
                    r.add_picture(doc_analyse.unpack_dir + "/word/" + path, c.img_size_x, c.img_size_y)
        #写入题目
        que_list = item[que]
        for que_item in que_list:
            for que_c_k in que_item:
                 aaa = que_item[que_c_k]
    doc.save(save_doc_path)  # 文档保存