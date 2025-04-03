# coding=UTF-8
import os
import sys
import docx
import re
from docx import Document
# import xlrd
# import xlwt
# from xlutils.copy import copy
import xlwt as xlwt
from enum import Enum

from docx.opc.oxml import qn
from docx.shared import Pt, RGBColor

from Tools import getNumPairs, getImportPath, isReadingMain, split_options, get_first_number_before_dot, get_max_number, \
    get_paragraph_shading


class State(Enum):
    invaid = 0 # 无效
    main = 1 # 题干
    que = 2 # 题目
    opt = 3 # 选项
    answer = 4 # 答案\分类\分数\解析
    finish = 5

class QuestionData:
    id = 0
    oldId = 0
    main = ""
    qus = "题目"
    opt = []
    ans = ""
    analysis = "解析"
    category = "知识点"
    qusType = "题型"
    mark = 1
    isSignle:bool = False

    def __init__(self, ques:str):

        self.id = int(ques.split(".")[0])
        self.oldId = self.id
        self.qus = ques
        self.opt = []
        self.ans = "答案"
        self.analysis = ""
        self.isSignle = False
        self.main = ""
        self._optStr = ""
        self._ansStr = ""

    def AddMain(self,mainContent):
        self.main = self.main + mainContent

    def AddQueContent(self, ques):
        self.qus = self.qus + ques

    def AddOption(self,opts):
        self._optStr = self._optStr + "\t" + opts

    def UpdateOpt(self):
        self.opt = split_options(self._optStr)
        # optstr = self._optStr.replace("\t", "")
        # optstr = optstr.replace("A.", "A.").replace("A .", "A.")
        # optstr = optstr.replace("B.", "#B.").replace("B .", "#B.")
        # optstr = optstr.replace("C.", "#C.").replace("C .", "#C.")
        # optstr = optstr.replace("D.", "#D.").replace("D .", "#D.")
        # optstr = optstr.replace("E.", "#E.").replace("E .", "#E.")
        # optstr = optstr.replace("F.", "#F.").replace("F .", "#F.")
        # optstr = optstr.strip()
        # arr = optstr.split("#")
        # # self.opt = self.opt + arr
        # for a in arr:
        #     if a != '' and len(a) > 0:
        #         self.opt.append(a)
        if len(self.opt) > 6:
            print("id:%d 选项格式有错(或者下一题的题目格式有错,题目一定是数字+.的格式)，请检查,选项超过4个 %s" % (self.oldId, " ".join(self.opt)))
            return False
        self.opt.sort()
        return True

    def AddAnswer(self, answer):
        self._ansStr = self._ansStr + "\n" + answer

    def updateAns(self):
        if self._ansStr.find("【答案】") != -1:

            answer = self._ansStr.replace("．",".")
            answer = answer.replace("     ", ";")
            answer = answer.replace("【答案】","")
            answer = answer.strip()
            tmpAnswer = answer
            if tmpAnswer.find("】") -1:
                tmpAnswer = tmpAnswer[:tmpAnswer.find("【")]
            # 使用正则表达式匹配数字和对应的答案
            if tmpAnswer.find(".") != -1:
                first_num:int = get_first_number_before_dot(tmpAnswer)
                max_num:int = get_max_number(tmpAnswer)
                answers = {}
                for i in range(max_num - first_num + 1):
                    tmp_id: int = first_num + i
                    r: str = getNumPairs(tmp_id, tmpAnswer, None, max_num)
                    answers[tmp_id] = r
                matches = re.findall(r'(\d+)\.\s*(\S+)', tmpAnswer)

                # if matches:
                #     for match in matches:
                #         number, a = match
                #         if a[0] == ";":
                #             a = a[1:]
                #         answers[number] = a

                if not answers.__contains__(self.id):
                    print("id:%d 答案找不到,当前答案是:%s" % (self.id,answer))
                    return
                self.ans = answers[self.id]
            else:
                self.ans = answer
        else:
            self.analysis = self.analysis + self._ansStr
        if len(self.ans) > 80 and self.qus.__contains__("____"):
            print("被识别为填空题，但是答案超过80的长度，建议改成 [**请回答**] 。 题目:%s" % self.qus)

    def updateAnl(self,max:int):
        paragraphs = []
        current_paragraph = []
        self.qus = self.qus.strip()
        self.qus = self.qus.replace("("," [").replace("（"," [").replace(")","] ").replace("）","] ")
        self._ansStr = self._ansStr.replace("【甲】","[甲]").replace("【乙】","[乙]").replace("【丙】","[丙]").replace("【丁】","[丁]")
        analysis1:str = self._ansStr.replace("【分析】","【解析】").replace("【详解】","【解析】")
        pattern = r'【解析】(.*?)(?=【)'
        if not self.isSignle and not analysis1.__contains__("【点睛】"):
            analysis1 = analysis1 + "【"
        matches = re.findall(pattern, analysis1, re.DOTALL)
        anlDic = {}
        if not self.isSignle and len(matches) == 0:
            print("id=%d 解析中答案查找失败" % self.id)
        for idx, item in enumerate(matches, start=1):
            text = item.strip()
            tmpAnl = getNumPairs(self.id,text,"【",max )
            if not tmpAnl:
                if self.isSignle :
                    return
                print("找不到解析: id = %d,解析如下:%s\n========================\n\n\n" % (self.id,self._ansStr) )
                return
            tmpAnl = "【解析】"+tmpAnl
            self.analysis = tmpAnl
            if self.analysis.endswith("【"):
                self.analysis = self.analysis[:-1]
            return

    def AddAnalysis(self, analysis):
        self.analysis = analysis

    def AddCategory(self,value):
        self.category = value

    def AddQusType(self,qusType):
        self.qusType = qusType

    def AddMark(self,value):
        self.mark = value

    def tostring(self):
        return str(self.id) + "." + self.qus + " 选项:" + " ".join(self.opt) + " 答案:" + self.ans;

class ComprehensionData:
    id:int = 0
    ques:[] = []
    main:str = "" # 题干

    def __init__(self):
        global queId
        queId = queId + 1
        self.id = queId
        self.ques = []
        self.main = ""


    def addMainContent(self,mainContent:str):
        mainContent = mainContent.replace("、",",")
        m = re.match( r'^#\d+\.',mainContent)
        if m :
            print("题干的内容不能以数字+点开头 :" + mainContent)
        self.main = self.main + mainContent

    def addQue(self,que:QuestionData):
        self.ques.append(que)
        tmpId = len(self.ques)
        que.id = tmpId


curState = State.invaid
queId = 0
isValid = False
comprehensionArr = []
curComprehensItem:ComprehensionData = None
curQuestionDataItem:QuestionData = None
last_color = None

def main():
    if len(sys.argv) > 1:
        fileName = sys.argv[1]
    else:
        fileName = getImportPath()
    if not fileName.endswith(".docx"): # 文件夹
        for path in os.listdir(fileName):
            if path.endswith(".docx"):
                path = fileName + "/" + path
                HandleFile(path)
    else:
        HandleFile(fileName)

def formatUnderline(d:docx.text.paragraph.Paragraph):
    return
    # for item in d.runs :
    #     if item.underline :
            # item.text = item.text.replace(" ","_")
            # item.text = "____________"

def replaceCommon(line:str):
    line = line.replace("【指出通假】","[指出通假]").replace("【借助工具书】","[借助工具书]")
    line = line.replace("【同类现象】","[同类现象]").replace("．", ".")
    line = line.replace("【链接材料一】","[链接材料一]").replace("【链接材料二】","[链接材料二]")
    line = line.replace("【示例一】", "[示例一]").replace("【示例二】", "[示例二]").replace("【示例三】", "[示例三]")

    line = line.replace("【链接材料】", "[链接材料]")
    line = line.replace("【分析】", "【解析】").replace("【详解】", "【解析】").replace( "【导语】","【解析】")
    line = line.replace("【整体分析】" ,"【解析】")
    line = line.replace("【甲】","[甲]").replace("【乙】","[乙]").replace("【丙】","[丙]").replace("【丁】","[丁]")
    for i in range(20):
        line = line.replace("【小题" + str(i+1)+"】",str(i+1)+".")
    return line

def HandleFile(fileName:str):
    global curState,queId,isValid,comprehensionArr,curComprehensItem,curQuestionDataItem
    curState = State.invaid
    queId = 0
    isValid = False
    comprehensionArr = []
    curComprehensItem = None
    curQuestionDataItem = None

    if fileName.__contains__("[上传]"):
        return
    if not os.path.exists(fileName):
        print( "文件不存在%s" % fileName)
        return
    doc = getDoc(fileName)
    print("----------------开始处理文件 %s ------------" % fileName)
    global category
    global type
    name, category, type, mark = AnalyzeFileName(fileName)

    global questionRule
    global optionROle
    # global  curState
    questionRule = re.compile(r"^\d+\s*\.\s*")
    optionROle = re.compile(r"^[A-Z]\s*\.\s*")
    for d in doc.paragraphs:
        formatUnderline(d)
        str = replaceCommon(d.text)
        if str == '':
            continue
        if not checkContent(str,d):
            continue
        if curState == State.finish :
            break
        if curState == State.invaid:
            continue
        if curState == State.main:
            str = "\n        " + str
        updateState(str)

    ## 写答案
    fileName = fileName.split(".")[0] + "[上传].docx"
    writeAns(fileName)
    print("----------------结束处理文件----------------------------\n")
    return True

def checkContent(s:str,d):
    global curState
    global isValid
    global curComprehensItem
    global  curQuestionDataItem
    global last_color
    if s.find("//题目开始") != -1:
        isValid = True
        curComprehensItem = None
        curQuestionDataItem = None
        curState = State.main
        return False
    if s.find("//题目结束") != -1:
        isValid = False
        curState = State.finish
        return False
    color = get_paragraph_shading(d)
    is_main_start = False
    if last_color and not color:
        is_main_start = True
    last_color = color
    if isValid and ( is_main_start or isReadingMain(s)) :
        s = s.strip()
        curComprehensItem = None
        curQuestionDataItem = None
        curState = State.main
        return True
    if isValid:
        # 只有题干或选项，才有可能是题目
        if ( curState == State.main or curState == State.opt ) and isQustion(s): # 题目
            curQuestionDataItem = None
            curState = State.que
            return True
        if ( curState == State.que or curState == State.opt ) and isOption(s):  # 选项
            curState = State.opt
            return True
        if s.find("【答案】") != -1 or s.find("【知识点】") != -1 or s.find("【解析】") != -1 :
            curState = State.answer
            return True
        return True
    return False

def updateState(s:str):
    global curState
    global isValid
    global comprehensionArr
    global curComprehensItem
    global curQuestionDataItem
    if not isValid :
        return;
    match(curState):
        case State.main:
            pattern = r'^#\d+\.'
            m = re.match(pattern, s.strip())
            if curComprehensItem is None:
                curComprehensItem = ComprehensionData()
                comprehensionArr.append(curComprehensItem)
            if s.find("#") != -1:
                if m:
                    mlen = len(m.group(0))
                    ss = s[mlen:]
                    curComprehensItem.addMainContent(ss)
                else:
                    curComprehensItem.addMainContent(s.split("#")[-1])
            else:
                curComprehensItem.addMainContent(s)
            # 如果题干和题目是一起
            if m:
                curQuestionDataItem = QuestionData(s.split("#")[-1])
                curQuestionDataItem.isSignle = True
                curComprehensItem.ques.append(curQuestionDataItem)
        case State.que:
            s = s.replace("、", ",")
            if curQuestionDataItem is None:
                curQuestionDataItem = QuestionData(s)
                curComprehensItem.ques.append(curQuestionDataItem)
            else:
                if s.__contains__("."):
                    #说明上一题只有题目没有答案
                    if s.split(".")[0].isdigit() and curQuestionDataItem.id + 1 == int(s.split(".")[0]):
                        curQuestionDataItem = QuestionData(s)
                        curComprehensItem.ques.append(curQuestionDataItem)
                        return
                curQuestionDataItem.AddQueContent(s)
        case State.opt:
            if curQuestionDataItem is None:
                print("没有题目，直接就是选项目，请检查题干为:%s " % curComprehensItem.main)
                return
            curQuestionDataItem.AddOption(s)
        case State.answer:
            for item in curComprehensItem.ques:
                item.AddAnswer(s)



def isOption(str):
    global optionROle
    if optionROle.match(str):
        return True
    return False

def AnalyzeFileName(str):
    pathArr = str.split("/")
    str = pathArr[len(pathArr)-1]
    titleNameArr = str.split(".")
    if len(titleNameArr) >= 1 :
        str = titleNameArr[0]
    arr = str.split("_")
    if len(arr) == 0 :
        return str,"","",1
    name = arr[0]
    category = ""
    type = ""
    mark = 1
    if len(arr) > 0 :
        category = arr[1]
    if len(arr) > 1 :
        type = arr[2]
    return name,category,type,mark

def isQustion(str):
    if len(str) <= 1 :
        return False
    global questionRule
    if questionRule.match(str):
        return True
    return False

def isAns(str):
    global  analysisRule
    if analysisRule.match(str):
        return True
    return False


def getDoc(filename):
    d = docx.Document(filename)
    return d

def writeAns(fileName):
    '''python-docx实现操作word文档基础命令(包含插入各级标题)
    :param fileName:文件保存路径
    :return: None
    '''

    global category
    global type

    content:str = ""
    comNum = 0
    allQue:int = 0
    contents:[] = []
    split_num:int = 3
    file_per: int = split_num
    if len(comprehensionArr) >= split_num:
        file_per = int(len(comprehensionArr) / split_num)
    else:
        split_num = 1
    for c in comprehensionArr:
        comNum += 1
        com:ComprehensionData = c
        if len(com.ques) == 0 :
            print("题目数量为0:题干:" + com.main)
        allQue += len(com.ques)
        if com.main == "" :
            if len(com.ques) > 1:
                print("无题干:题目" + com.ques[0].qus)
            else:
                com.main = com.ques[0].qus
        com.main = com.main.replace("【","\n【")
        com.main = com.main.replace("("," [").replace("（"," [").replace(")","] ").replace("）","] ")
        if bool(re.match(r'^\d+\.', com.main.strip())):
            print("题干中，不能以 \"数字+.\" 开头，如\"1.\" ，请将\".\"改成其他字符 , 题干 =  %s " % com.main )
        content = content.rstrip()
        content += "\n[理解题开始]\n"
        com.main = com.main + "\n分类:"+type
        content += ( com.main.lstrip() + " (" + str(len(com.ques)) + "题)\n\n\n" )
        num = 0
        for q in com.ques:
            que:QuestionData = q
            que.UpdateOpt()
            que.updateAns()
            que.updateAnl( com.ques[0].id + len(com.ques) - 1)
            if que.ans == "":
                print("答案为空,id=%d" % que.id )
            # 题目
            content += que.qus + "\n"
            for op in que.opt:
                if len(op.strip()) <= 2:
                    print("选项为空，请注意检查 id = %d,选项:%s" % (que.id,op) )
                content += op + "\n"
            content += "答案:" + que.ans
            content = content.rstrip() + "\n"
            content += "分数:5\n"
            que.analysis = que.analysis.replace("【","\n【")
            num += 1
            if num == len(com.ques):
                content += "解析:" + que.analysis + "\n\n"
            else:
                content += "解析:" + que.analysis + "\n\n\n"
        content = content.rstrip()
        content += "\n[理解题结束]"
        if len(contents) < split_num - 1 and split_num > 1 and comNum == file_per :
            cs = {}
            cs["content"] = content
            cs["num"] = comNum
            contents.append(cs)
            content = ""
            comNum = 0
    if len(contents) < split_num :
        cs = {}
        cs["content"] = content
        cs["num"] = comNum
        contents.append(cs)
    file_name:int = 0
    for cs in contents:
        file_name += 1
        c = cs["content"]
        c = c.rstrip() + "\r\r"
        num = cs["num"]
        c = c.lstrip()
        doc = Document()  # 创建新文档
        para1 = doc.add_paragraph()
        run_2 = para1.add_run(c)  # 以add_run的方式追加内容，方便后续格式调整
        run_2.font.name = 'Times New Roman'  # 注：这个好像设置 run 中的西文字体
        # 设置中文字体
        # 需导入 qn 模块
        from docx.oxml.ns import qn
        # run_2.font.name = '楷体'  # 注：如果想要设置中文字体，需在前面加上这一句
        run_2.font.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
        # 设置字体大小
        run_2.font.size = Pt(14)
        #run.font.size = Pt(10.5)
        arr:[] = fileName.split(".")
        f_file_name:str = arr[0] + str(file_name) + "." + arr[1]
        if os.path.exists(f_file_name):
            os.remove(f_file_name)
        doc.save(f_file_name)  # 文档保存
        print("保存完成,题干数量%d,子题目数量%d" % (num, allQue) )



if __name__ == '__main__':
    main()
