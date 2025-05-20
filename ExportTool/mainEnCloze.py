# coding=UTF-8
import os
import sys
import docx
import re
from docx import Document
from enum import Enum

from docx.shared import Pt, RGBColor

from Tools import getNumPairs, getImportPath, split_options, getContentBeforeNumAndPot, get_first_number_before_dot, \
    get_max_number, is_start_with_num_and_point, get_paragraph_shading


class State(Enum):
    invaid = 0  # 无效
    main = 1  # 题干
    opt = 3  # 选项
    answer = 4  # 答案
    origin = 5  # 来源
    difficulty = 6  # 难度
    point = 7  # 知识点
    anly = 8  # 解析或导语
    dian_jing = 9 # 点晴
    finish = 100


class BaseItem:

    def __init__(self):
        self.content = ""

    def addContent(self, add_str: str):
        self.content = self.content + add_str

    #解析
    def update(self):
        return


#选项
class OptionItem(BaseItem):

    def __init__(self):
        super().__init__()
        self.opt_id = 0
        self.opt_arr = []

    def update(self):
        opt_id_str:str = self.content.split(".")[0]
        self.opt_id = int(opt_id_str)
        self.opt_arr = split_options(self.content)

    def getResult(self,score:int ):
        # r:str = "\n%d." % self.opt_id
        r: str = ""
        isFirst:bool = True
        for a in self.opt_arr:
            r = r + "\n" + str(a).strip()
            if isFirst:
                r = r + " =>" + str(score)
            isFirst = False
        return r

#知识点
class KnowledgePointItem(BaseItem):
    def __init__(self):
        super().__init__()
        self.desc = ""


#答案
class AnswerItem(BaseItem):
    def __init__(self):
        super().__init__()
        self._answer_str = None

    def update(self):
        self._answer_str = ""
        from_id:int = get_first_number_before_dot(self.content)
        to_id:int = get_max_number(self.content)
        for i in range(to_id-from_id+1):
            tmp_id:int = from_id + i
            r:str = getNumPairs(tmp_id, self.content, None, to_id)
            self._answer_str = self._answer_str + r.strip()

    def getResult(self):
        return self._answer_str



#解析
class AnalysisItem(BaseItem):

    def __init__(self):
        super().__init__()
        self.desc = None
        self.dian_jing:str = ""
        self.anlyItems = []

    def update(self):
        super().update()
        self.desc = getContentBeforeNumAndPot(self.content)
        first_id: int = get_first_number_before_dot(self.content)
        if first_id is None:
            print("解析中没有题号,解析内容:%s" % self.content)
            return
        if first_id == 0 :
            print("解析中有疑似有0.几的数字，可以尝试修改数字")
        if self.content.__contains__("【点睛】"):
            arr = self.content.split("【点睛】")
            if len(arr) > 1:
                self.dian_jing = "【点睛】" + arr[1]
                self.dian_jing = self.dian_jing.replace("\n\n","\n")
            self.content = arr[0]

        max_id: int = get_max_number(self.content)
        if max_id is None:
            print("解析序号最大值错误:first_id=%d,content=%s" % (first_id ,self.content))
            return
        for i in range(max_id - first_id + 1):
            anly = getNumPairs(first_id + i, self.content)
            if anly is None or anly == "":
                print("解析内容为空:%s" % self.content)
                return
            self.anlyItems.append( str(first_id + i) + ":" + anly)

    def getResult(self):
        return "".join(self.anlyItems) + self.dian_jing


class MainItem(BaseItem):

    def __init__(self):
        super().__init__()
        # 选项数组
        self.options = []
        # 知识点
        self.main_knowledge_point: KnowledgePointItem = None
        # 答案
        self.answer = None
        # 解析
        self.analysis: AnalysisItem = None
        # 来源
        self.origin = None
        # 难度
        self.difficulty = None

    def update(self):
        super().update()
        self.replacePot()
        for o in self.options:
            o.update()
        if self.main_knowledge_point:
            self.main_knowledge_point.update()
        if self.answer is None:
            print("请确定井号位置，题干为:%s" % self.content )
        else:
            self.answer.update()
        if self.analysis :
            self.analysis.update()
        self.content = self.content.replace("A、","A和").replace("B、","B和").replace("C、","C和")
        if not self.content.__contains__("] ( ) ") and not self.content.__contains__("__"):
            print("题目里居然没有任何选项: %s" % self.content)

    def replacePot(self):
        arr = ["A","B","C","D","E","F"]
        for a in arr:
            tmp:str = " " + a + "."
            if self.content.__contains__(tmp):
                new_str:str = " " + a + " ."
                self.content = self.content.replace( tmp , new_str)


    def getResult(self,id:int):
        r:str = "%d.\n" % id
        if self.difficulty:
            r = r + self.difficulty.content
        if self.origin:
            r = r + self.origin.content.replace("（","<").replace("）",">").replace("(","<").replace(")",">")
        if self.main_knowledge_point:
            r = r + "\n" + self.main_knowledge_point.content.replace("（","<").replace("）",">").replace("(","<").replace(")",">")
        self.content = self.content.replace("#","")
        self.content = self.content.lstrip()
        r = r.rstrip()
        r = r + "\n" + self.content
        score:int = 2
        for o in self.options:
            r = r + o.getResult(score)
        if self.answer :
            r = r.rstrip()
            r = r + "\n答案:" + self.answer.getResult() + "\n"
        r = r + "分数:" + str(score*len(self.options)) + "\n"
        r = r + "分类:完形填空\n"
        # r = r + "标签:" + self.main_knowledge_point.content.replace("【知识点】","") + "\n"
        if self.analysis is None:
            print("无解析: %s " % self.content )
        else:
            r = r +  "解析:" + self.analysis.getResult() + "\n\n\n"
            if len(self.options) != len(self.analysis.anlyItems):
                if len(self.analysis.anlyItems) > 0 :
                    all_anly = "\n".join(self.analysis.anlyItems)
                    print("选项和解析数量不匹配, id = %d , 选项:%d个 , 解析:%d ,  题干:%s \n======= 当前解析的内容======\n %s" % (
                    id, len(self.options), len(self.analysis.anlyItems), self.content, all_anly))
                else:
                    print("选项和解析数量不匹配, id = %d , 选项:%d个 , 解析:%d ,  题干:%s" % (id, len(self.options) ,len(self.analysis.anlyItems), self.content ))


        return r

    def add_option(self, opt: OptionItem):
        self.options.append(opt)

    def add_answer(self, answer_item: AnswerItem):
        self.answer = answer_item

    def add_knowledge_point(self, item: KnowledgePointItem):
        self.main_knowledge_point = item

    def set_analysis(self, item: AnalysisItem):
        self.analysis = item


cur_state = State.invaid
is_valid = False
question_rule = re.compile(r"^\d+\s*\.\s*")
optionROle = re.compile(r"^[A-Z]\s*\.\s*")
mainItems = []
cur_main_item: MainItem = None
#当前选项
cur_opt_item: OptionItem = None
#答案
cur_answer_item: AnswerItem = None
#知识点
cur_knowledge_point: KnowledgePointItem = None
#解析
cur_analysis_item: AnalysisItem = None
#难度
cur_difficulty_item:BaseItem = None
#来源
cur_original_item:BaseItem = None


def main():
    if len(sys.argv) > 1:
        fileName = sys.argv[1]
    else:
        fileName = getImportPath()
    if not fileName.endswith(".docx"):  # 文件夹
        for path in os.listdir(fileName):
            if path.endswith(".docx"):
                path = fileName + "/" + path
                HandleFile(path)
    else:
        HandleFile(fileName)


def common_repalce(line:str):
    line = line.replace("【导语】","【解析】").replace("【导读】","【解析】").replace("【分析】","【解析】").replace("A.M.","AM").replace("A.D.","AD").replace("P.M.","PM")
    line = line.replace("【详解】","【解析】")
    line = line.replace("$" ,"$ ")
    return line

def formatUnderline(d:docx.text.paragraph.Paragraph):
    for item in d.runs :
        if item.underline :
            s:str = item.text.strip()
            if s.isdigit():

                item.text = "[%s] ( ) " % s
                item.underline = False
                #print(item.text)
                # item.text = item.text.replace(" ","_")
                # item.text = "____________"

def HandleFile(file_name: str):
    global cur_state, is_valid, mainItems, cur_main_item, cur_opt_item, cur_answer_item, cur_knowledge_point, cur_analysis_item
    cur_state = State.invaid
    is_valid = False
    mainItems = []
    cur_main_item = None
    cur_opt_item = None
    cur_answer_item = None
    cur_knowledge_point = None
    cur_analysis_item = None
    cur_original_item = None
    cur_difficulty_item = None
    if file_name.__contains__("[上传]"):
        print("选错文件了，里面包含了[上传]")
        return
    if not os.path.exists(file_name):
        print("文件不存在%s" % file_name)
        return
    doc = getDoc(file_name)
    print("----------------开始处理文件 %s ------------" % file_name)
    global category
    global type,is_anly_over
    #上一行是不是解析结尾
    is_last_anly_over:bool = False
    name, category, type, mark = AnalyzeFileName(file_name)
    is_last_start:bool = False
    first_num = None
    is_last_anly = False
    last_color = None
    for d in doc.paragraphs:
        formatUnderline(d)
        line = d.text.replace("．", ".")
        if line == '':
            continue
        if not line.startswith("【"):
            line = "\n        "+line
        line = common_repalce(line)
        color = get_paragraph_shading(d)
        is_new_color = False
        if color is None and last_color:
            is_new_color = True
        last_color = color
        is_new_item = checkState(line)
        if cur_state == State.finish:
            break
        if cur_state == State.invaid:
            if not is_last_start:
                is_last_start = line.__contains__("//题目开始")
                if is_last_start:
                    cur_state = State.main
                    cur_main_item = None
                continue
            else:
                cur_state = State.main
        #如果没有识别出主题，但是上一次
        # first_num = is_start_with_num_and_point(line)
        # if not is_new_item and ( cur_state == State.anly or cur_state == State.dian_jing ) and not first_num and is_last_anly:
        #     cur_state = State.main
        #     cur_main_item = None
        if is_new_color:
            cur_state = State.main
            cur_main_item = None
        update_state(line, is_new_item)

        is_last_anly = is_endwith_anly(line)

    for item in mainItems:
        main_item: MainItem = item
        main_item.update()

    ## 写答案
    file_name = file_name.split(".")[0] + "[上传].docx"
    writeAns(file_name)
    print("----------------结束处理文件----------------------------\n")
    return True


def checkState(s: str):
    global cur_state
    global is_valid
    global is_anly_over
    is_new_item: bool = True
    s = s.strip()
    if s.find("//题目开始") != -1:
        is_valid = True
        return is_new_item
    if s.find("//题目结束") != -1:
        cur_state = State.finish
        return is_new_item
    if not is_valid:
        return
    if s.startswith("#"):
        cur_state = State.main
        return is_new_item
    # 只有题干或选项，才有可能是题目
    if cur_state.value <= State.opt.value and is_option(s):  # 选项
        cur_state = State.opt
        return is_new_item
    if cur_state.value <= State.answer.value and s.find("【答案】") != -1:
        cur_state = State.answer
        return is_new_item
    if cur_state.value <= State.point.value and s.find("【知识点】") != -1:
        cur_state = State.point
        return is_new_item
    if cur_state.value <= State.anly.value and s.find("【解析】") != -1:
        cur_state = State.anly
        return is_new_item
    if s.find("【点睛】") != -1:
        cur_state = State.dian_jing
        return is_new_item
    if s.find("【难度】") != -1:
        cur_state = State.difficulty
        return is_new_item
    if s.find("【来源】") != -1:
        cur_state = State.origin
        return is_new_item
    is_new_item = False
    return is_new_item

const_opt = ["A","B","C","D","E","F"]

def is_endwith_anly(s:str):
    if s.__contains__("故答案为") or s.__contains__("故选") or s.__contains__("所以答案") :
        return True
    for cs in const_opt:
        if s.__contains__("故"+cs + "项"):
            return True
        if s.__contains__("故选"+cs):
            return True
        if s.__contains__("选"+cs+ "项"):
            return True
        if s.__contains__("故此选"+cs):
            return True
        if s.__contains__("故" + cs + "项正确"):
            return True
        if s.__contains__("故" + cs + "选项正确"):
            return True
        if s.__contains__("答案应是" + cs):
            return True
        if s.__contains__("可知答案为" + cs):
            return True
        if s.__contains__("故此题选择" + cs):
            return True
        if s.__contains__("可知" + cs + "选项正确"):
            return True
    return False

def update_state(line: str, is_new: bool):
    global cur_state
    global is_valid, cur_main_item, mainItems, cur_opt_item, cur_answer_item, cur_knowledge_point, \
        cur_analysis_item,cur_difficulty_item,cur_original_item

    if not is_valid:
        return
    match cur_state:
        case State.main:  #主题文章
            if cur_main_item is None or is_new:
                cur_main_item = MainItem()
                mainItems.append(cur_main_item)
            cur_main_item.addContent(line)
        case State.opt:  #选项
            if cur_main_item is None:
                print("题干里面是不是没有. 内容: %s" % line)
                return
            if cur_opt_item is None or is_new:
                cur_opt_item = OptionItem()
                cur_main_item.add_option(cur_opt_item)
            cur_opt_item.addContent(line)
        case State.answer:  #答案
            if cur_main_item == None:
                print("题干里面是不是没有 #号")
            if is_new:
                cur_answer_item = AnswerItem()
                cur_main_item.add_answer(cur_answer_item)
            cur_answer_item.addContent(line)
        case State.point:
            if is_new:
                cur_knowledge_point = KnowledgePointItem()
                cur_main_item.add_knowledge_point(cur_knowledge_point)
            cur_knowledge_point.addContent(line)
        case State.anly:
            if is_new:
                cur_analysis_item = AnalysisItem()
                cur_main_item.set_analysis(cur_analysis_item)
            cur_analysis_item.addContent(line)
        case State.dian_jing:
            if not cur_analysis_item:
                cur_analysis_item = AnalysisItem()
                cur_main_item.set_analysis(cur_analysis_item)
            cur_analysis_item.addContent("\n"+line)
        case State.difficulty:
            if is_new:
                cur_difficulty_item = BaseItem()
                cur_main_item.difficulty = cur_difficulty_item
            cur_difficulty_item.addContent(line)
        case State.origin:
            if is_new:
                cur_original_item = BaseItem()
                cur_main_item.origin = cur_original_item
            cur_original_item.addContent(line)


def is_option(line: str):
    global question_rule
    if question_rule.match(line):
        return True
    return False


def AnalyzeFileName(str):
    pathArr = str.split("/")
    str = pathArr[len(pathArr) - 1]
    titleNameArr = str.split(".")
    if len(titleNameArr) >= 1:
        str = titleNameArr[0]
    arr = str.split("_")
    if len(arr) == 0:
        return str, "", "", 1
    name = arr[0]
    category = ""
    type = ""
    mark = 1
    if len(arr) > 0:
        category = arr[1]
    if len(arr) > 1:
        type = arr[2]
    return name, category, type, mark


def isQustion(str):
    if len(str) <= 1:
        return False
    global question_rule
    if question_rule.match(str):
        return True
    return False


def isAns(str):
    global analysisRule
    if analysisRule.match(str):
        return True
    return False


def getDoc(filename):
    d = docx.Document(filename)
    return d


def writeAns(fileName):
    global category

    global type
    global mainItems
    doc = Document()  # 创建新文档
    para1 = doc.add_paragraph()
    content: str = ""
    last_que_id:int = 0
    allQue: int = 0
    main_id:int = 1
    for main_item in mainItems:
        r = main_item.getResult(main_id)
        content = content + r + "\n\n"
        main_id = main_id + 1
        allQue = allQue + len(main_item.options)
        for o in main_item.options:
            if last_que_id + 1 != o.opt_id:
                print("子题序号不连续: 上一题号:%d,当前题号:%d，如果题号是连续的，请确认是不是有类似【点睛】的内容无法识别结束" % (last_que_id,o.opt_id) )
            last_que_id = o.opt_id

    run_2 = para1.add_run(content)  # 以add_run的方式追加内容，方便后续格式调整
    run_2.font.name = 'Times New Roman'  # 注：这个好像设置 run 中的西文字体
    # 设置中文字体
    # 需导入 qn 模块
    from docx.oxml.ns import qn
    # run_2.font.name = '楷体'  # 注：如果想要设置中文字体，需在前面加上这一句
    run_2.font.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
    # 设置字体大小
    run_2.font.size = Pt(14)
    if os.path.exists(fileName):
        os.remove(fileName)
    doc.save(fileName)  # 文档保存
    print("保存完成,题干数量%d,子题目数量%d" % (len(mainItems), allQue))


if __name__ == '__main__':
    main()

"""
1.优考试在线考试系统于        上线，是一款安全、稳定、专业的在线考试系统，支持SaaS在线考试和        ，考试全程智能化处理，比如考前        、考中技术支持、考后数据分析，客服全流程跟进服务，保障考试顺利流畅进行。且优考试拥有多达        种防作弊功能，为考试安全保驾护航。哈哈哈哈         OK  。 航。哈哈哈哈        OK 
1.
A. 2014
B.2015 => 2
C.2016
D.2017
2.
A. 局域网考试 => 2
B.线下考试
C.无网络考试
D. 飞书考试
3.
A. 压力测试 => 2
B.布置作业
C.培训
D. 观看视频
4.
A. 8
B.9
C.13 => 2
D. 10
5.
A. 8
B.9
C.13 => 2
D. 10
6.
A. 8
B.9
C.13 => 2
D. 10
答案：BAACDD
解析: 术支持、考后数据
"""
