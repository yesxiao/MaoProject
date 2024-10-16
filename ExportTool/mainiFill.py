# coding=UTF-8
import os
import sys
import docx
import re
from docx.shared import Inches
# import xlrd
# import xlwt
# from xlutils.copy import copy
import xlwt as xlwt

from Tools import getImportPath


class QuestionData:
    id = 0
    oldId = 0
    qus = "题目"
    opt = []
    ans = "答案"
    analysis = "解析"
    category = "分类"
    qusType = "题型"
    mark = 1

    def __init__(self, ques):
        quesarr = ques.split(".")
        self.id = int(quesarr[0])
        self.oldId = self.id
        del quesarr[0]
        self.qus = "".join(quesarr)
        if ques.endswith(".") :
            self.qus = self.qus + "."
        # self.qus = ques.replace(str(self.id)+".","")
        self.opt = []
        self.ans = "答案"
        self.analysis = "解析"

    def AddQueContent(self, ques):
        self.qus = self.qus + ques

    def AddOption(self, opts):
        optstr = opts.replace("\t", "")
        optstr = optstr.replace("A.", "A.").replace("A .", "A.")
        optstr = optstr.replace("B.", "#B.").replace("B .", "#B.")
        optstr = optstr.replace("C.", "#C.").replace("C .", "#C.")
        optstr = optstr.replace("D.", "#D.").replace("D .", "#D.")
        optstr = optstr.replace("E.", "#E.").replace("E .", "#E.")
        optstr = optstr.replace("F.", "#F.").replace("F .", "#F.")
        optstr = optstr.strip()
        arr = optstr.split("#")
        # self.opt = self.opt + arr
        for a in arr:
            if a != '' and len(a) > 0:
                self.opt.append(a)
        if len(self.opt) > 6:
            print("id:%d 选项格式有错(或者下一题的题目格式有错,题目一定是数字+.的格式)，请检查,选项超过4个 %s" % (self.oldId, " ".join(self.opt)))
            return False
        self.opt.sort()
        return True

    def AddAnswer(self, answer):
        self.ans = answer

    def AddAnalysis(self, analysis):
        self.analysis = analysis

    def AddCategory(self,value):
        self.category = value

    def AddQusType(self,qusType):
        self.qusType = qusType

    def AddMark(self,value):
        self.mark = value

    def tostring(self):
        return str(self.id) + "." + self.qus + " 选项:" + " ".join(self.opt) + " 答案:" + self.ans;


def main():
    if len(sys.argv) > 1:
        fileName = sys.argv[1]
    else:
        fileName = getImportPath()
    if not fileName.endswith(".docx"): # 文件夹
        for path in os.listdir(fileName):
            if path.startswith("~"):
                continue
            if path.endswith(".docx"):
                path = fileName + "/" + path
                HandleFile(path)
    else:
        HandleFile(fileName)

def formatUnderline(d:docx.text.paragraph.Paragraph):
    for item in d.runs :
        if item.underline :
            item.text = "____"
            # item.text = item.text.replace(" ","_")

def HandleFile(fileName:str):
    doc = getDoc(fileName)

    name, category, type, mark = AnalyzeFileName(fileName)

    global questionRule
    global optionROle
    global analysisRule
    # questionRule = re.compile("^[0-9]*\.")
    questionRule = re.compile("^\d+\s*\.\s*")
    # optionROle = re.compile("^[A-Z]*\.")
    optionROle = re.compile("^[A-Z]\s*\.\s*")
    analysisRule = re.compile("^【详解")
    questions = []
    questionsDic = {}
    isStartChoise = False
    isStartQus = False
    hasError = False
    curQu = None
    for d in doc.paragraphs:
        formatUnderline(d)
        str = d.text.replace("．", ".").replace("\t", "")
        if str == '':
            continue
        if str.find("//题目开始") != -1:
            isStartChoise = True
        elif str.find("//题目结束") != -1:
            isStartChoise = False
            break
        if not isStartChoise:
            continue
        if isQustion(str):  # 问题
            curQu = QuestionData(str)
            #if curQu.id != len(questions) + 1 :
            #    print("%s 题目 id = %d 未按顺序,或之前的" % (fileName, curQu.id))
            #curQu.id = len(questions) + 1
            curQu.AddCategory(category)
            curQu.AddQusType(type)
            curQu.AddMark(mark)
            questions.append(curQu)
            if questionsDic.__contains__(curQu.id):
                print("%s 题目 %d 存在重复" % (fileName, curQu.id))
                hasError = True
                break
            questionsDic[curQu.id] = curQu
            isStartQus = True
        elif not isOption(str) and isStartQus:
            curQu.AddQueContent(str)
        elif isOption(str):  # 选项
            if not curQu:
                print("选项 %s 没有对应的题目" % str)
                hasError = True
                break
            if not curQu.AddOption(str):
                hasError = True
                break

    ## 写答案
    answner = False
    isStartAns = False
    analysisItem = None
    lastItem = None
    answnerDic = {}
    for i in range(0,len(doc.paragraphs)):
        d = doc.paragraphs[i]
        str = d.text.replace("．", ".").replace("\t", "")
        if str.find("//答案开始") != -1:
            answner = True
        elif str.find("//答案结束") != -1:
            if analysisItem and lastItem:
                lastItem.AddAnalysis(analysisItem)
                answnerDic[lastItem.id] = lastItem
                lastItem = None
            break
        if not answner and not isStartAns:
            continue
        if isQustion(str):  # 问题 和 答案 一样
            if analysisItem and lastItem:
                lastItem.AddAnalysis(analysisItem)
                lastItem = None
            arr = str.split(".")
            tmpId = int(arr[0])
            tmpAns = arr[1]
            if not questionsDic.__contains__(tmpId):
                print("%s %d 题目不存在" % (fileName,tmpId))
                hasError = True
                continue
                # break
            if answnerDic.__contains__(tmpId):
                print("%s %d 答案重复，请确定答案是否重复，或答案中有类似 %d.xxx 的格式，此格式会影响答案识别 . line: %s" % (fileName,tmpId, tmpId,d.text))
                hasError = True
                # break
                continue
            tmpQu = questionsDic[tmpId]
            answnerDic[tmpId] = tmpQu
            tmpQu.AddAnswer(tmpAns)
            lastItem = tmpQu
            isStartAns = True
            analysisItem = []
        elif isStartAns and answner and isAns(str):
            analysisItem.append(str)

    if not hasError:
        fileName = fileName.split(".")[0] + ".xlsx"
        writeToExcel(fileName, "题目", questions)
        return True
    return False

def isOption(str):
    global optionROle
    if optionROle.match(str):
        return True
    return False

def AnalyzeFileName(str):
    pathArr = str.split("/")
    str = pathArr[len(pathArr)-1]
    titleNameArr = str.split(".")
    if len(titleNameArr) >= 1 :
        str = titleNameArr[0]
    arr = str.split("_")
    if len(arr) == 0 :
        return str,"","",1
    name = arr[0]
    category = ""
    type = ""
    mark = 1
    if len(arr) > 1 :
        category = arr[1]
    if len(arr) > 2 :
        type = arr[2]
    return name,category,type,mark

def isQustion(str):
    global questionRule
    if questionRule.match(str):
        return True
    return False

def isAns(str):
    global  analysisRule
    if analysisRule.match(str):
        return True
    return False


def getDoc(filename):
    d = docx.Document(filename)
    return d

def setColWidth( sheet , colIdx , width ):
    sheet.col(colIdx).width = width

def setRowHeight( sheet,rowIdx ,height ):
    alignment = xlwt.Alignment()
    alignment.horz = xlwt.Alignment.HORZ_CENTER
    alignment.vert = xlwt.Alignment.VERT_CENTER
    contentSytle = xlwt.easyxf('font:height %d ' % height)
    contentSytle.alignment = alignment
    sheet.row(rowIdx).set_style(contentSytle)

def setAutoline( style , isAutoline ):
    ##style.alignment.wrap = isAutoline
    if isAutoline == 1 :
        newlineStyle = xlwt.easyxf('font:height 250')
        alignment = xlwt.Alignment()
        alignment.horz = xlwt.Alignment.HORZ_LEFT
        alignment.vert = xlwt.Alignment.VERT_CENTER
        alignment.wrap = 1
        newlineStyle.alignment = alignment
        return newlineStyle
    return style


def writeToExcel(path, sheet_name, questionDatas):
    index = len(questionDatas)  ##读取所需要的行
    workbook = xlwt.Workbook()
    sheet:xlwt.Worksheet = workbook.add_sheet(sheet_name)
    alignment = xlwt.Alignment()
    alignment.horz = xlwt.Alignment.HORZ_LEFT
    alignment.vert = xlwt.Alignment.VERT_CENTER
    titleSytle = xlwt.easyxf('font:height 250')
    titleSytle.alignment = alignment
    lineStyle = xlwt.easyxf('font:height 250')
    lineStyle.alignment = alignment



    titles = ["序号","题目（必填）","题型(必填)",## "选项A（必填）","选项B（必填）","选项C（必填）","选项D（必填）","选项E","选项F",
               "正确答案1（必填）","正确答案2","正确答案3","正确答案4","正确答案5","正确答案6","解析（必填）","分数（必填）","标签"]
    sizes = [100,   1000,           400,        400,            400,         400,           400,        100,    100,    100,            300,       100 ,       100]
    autolines = [0,     1,             0,          1,               0,           0,            0,          0,     0,      0,              0,         0,           0]
    for i in range(0,len(titles)):
        sheet.write(0, i, titles[i],titleSytle)
        if i < len(sizes):
            setColWidth(sheet,i,sizes[i]*20)
    j = 0
    setRowHeight(sheet, 0, 400)
    for i in range(0, index):
        que = questionDatas[i]
        row = i + 1

        sheet.write(row, j, que.id,setAutoline(lineStyle,autolines[j]))
        j = j + 1
        sheet.write(row, j, que.qus,setAutoline(lineStyle,autolines[j]))
        j = j + 1
        sheet.write(row,j,que.qusType,setAutoline(lineStyle,autolines[j]))
        j = j + 1
        # optLen = len(que.opt)
        # if optLen < 6:
        #     optLen = 6
        # for n in range(j, j + len(que.opt)):
        #     tmpOp = que.opt[n - j]
        #     tmpOp = tmpOp.replace("A.", "").replace("B.", "").replace("C.", "").replace("D.", "").replace("E.","").replace("F.","")
        #     sheet.write(row, n, tmpOp,setAutoline(lineStyle,autolines[j]))
        # j = j + optLen
        arr = que.ans.split(",")
        for n in range(0,6):
            if n < len(arr):
                sheet.write(row, j, arr[n], setAutoline(lineStyle, autolines[j]))
            j = j + 1
        # sheet.write(row, j, que.ans,setAutoline(lineStyle,autolines[j]))
        #j = j + 1
        sheet.write(row, j , que.analysis,setAutoline(lineStyle,autolines[j]))
        j = j + 1
        sheet.write(row,j,que.mark,setAutoline(lineStyle,autolines[j]))
        j = j + 1
        sheet.write(row,j,que.category,setAutoline(lineStyle,autolines[j]))
        j = 0
        setRowHeight(sheet, i, 400)

    workbook.save(path)


if __name__ == '__main__':
    main()
