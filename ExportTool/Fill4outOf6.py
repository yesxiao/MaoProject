# coding=UTF-8
import os
import sys
import docx
import re
from docx import Document
# import xlrd
# import xlwt
# from xlutils.copy import copy
import xlwt as xlwt
import xml.etree.ElementTree as ET
from enum import Enum

from docx.opc.oxml import qn
from docx.shared import Pt, RGBColor

from Tools import getNumPairs, getImportPath, split_options, get_paragraph_shading, getContentBeforeNumAndPot, \
    get_max_number, get_first_number_before_dot


class State(Enum):
    invaid = 0 # 无效
    main = 1 # 题干
    que = 2 # 题目
    opt = 3 # 选项
    answer = 4 # 答案\分类\分数\解析
    analyse = 5
    come_from = 6
    finish = 7

class QuestionData:
    id = 0
    oldId = 0
    main = ""
    qus = "题目"
    opt = []
    ans = ""
    analysis = "解析"
    category = "知识点"
    qusType = "题型"
    mark = 1
    isSignle:bool = False

    def __init__(self, ques:str):

        self.id = int(ques.split(".")[0])
        self.oldId = self.id
        self.qus = ques
        self.opt = []
        self.ans = "答案"
        self.analysis = ""
        self.isSignle = False
        self.main = ""
        self._optStr = ""

    def AddMain(self,mainContent):
        self.main = self.main + mainContent

    def AddQueContent(self, ques):
        self.qus = self.qus + ques

    def AddOption(self, opts):
        self._optStr = self._optStr + "\t" + opts
        return True

    def UpdateOpt(self):
        self.opt = split_options(self._optStr)
        if len(self.opt) > 6:
            print("id:%d 选项格式有错(或者下一题的题目格式有错,题目一定是数字+.的格式)，请检查,选项超过4个 %s" % (self.oldId, " ".join(self.opt)))
            return False
        self.opt.sort()
        return True

    def AddAnswer(self, answer):
        if answer.find("【答案】") != -1:
            answer = answer.replace("．",".")
            answer = answer.replace("【答案】","")
            answer = answer.strip()
            # 使用正则表达式匹配数字和对应的答案
            if answer.find(".") != -1:
                matches = re.findall(r'(\d+)\.\s*(\S+)', answer)
                answers = {}
                if matches:
                    for match in matches:
                        number, a = match
                        answers[number] = a

                if not answers.__contains__(str(self.id)):
                    print("id:%d 答案找不到,当前答案是:%s" % (self.id,answer))
                    return
                self.ans = answers[str(self.id)]
            else:
                self.ans = answer
        else:
            self.analysis = self.analysis + answer

    def updateAnl(self,max:int):
        paragraphs = []
        current_paragraph = []
        self.qus = self.qus.replace("("," [").replace("（"," [").replace(")","] ").replace("）","] ")
        self.analysis = self.analysis.replace("【分析】","【解析】").replace("【详解】","【解析】")
        pattern = r'【解析】(.*?)(?=【)'
        if not self.isSignle and not self.analysis.__contains__("【点睛】"):
            self.analysis = self.analysis + "【"
        matches = re.findall(pattern, self.analysis, re.DOTALL)
        anlDic = {}
        if not self.isSignle and len(matches) == 0:
            print("id=%d 解析中答案查找失败" % self.id)
        for idx, item in enumerate(matches, start=1):
            text = item.strip()
            tmpAnl = getNumPairs(self.id,text,"【",max )
            if not tmpAnl:
                if self.isSignle :
                    return
                print("找不到解析: id = %d" % self.id )
                return
            tmpAnl = re.sub(pattern,"【解析】"+tmpAnl,self.analysis,flags=re.DOTALL)
            self.analysis = tmpAnl
            if self.analysis.endswith("【"):
                self.analysis = self.analysis[:-1]
            return

    def AddAnalysis(self, analysis):
        self.analysis = analysis

    def AddCategory(self,value):
        self.category = value

    def AddQusType(self,qusType):
        self.qusType = qusType

    def AddMark(self,value):
        self.mark = value

    def tostring(self):
        return str(self.id) + "." + self.qus + " 选项:" + " ".join(self.opt) + " 答案:" + self.ans;

class ComprehensionData:
    id:int = 0
    ques:[] = []
    main:str = "" # 题干

    def __init__(self):
        global queId
        queId = queId + 1
        self.id = queId
        self.ques = []
        self.main = ""
        self.answer = ""
        self.analyse = ""
        self.come = ""


    def addMainContent(self,mainContent:str):
        mainContent = mainContent.replace("、",",")
        m = re.match( r'^#\d+\.',mainContent)
        if m :
            print("题干的内容不能以数字+点开头 :" + mainContent)
        self.main = self.main + mainContent

    def add_answer(self,s:str):
        self.answer = self.answer + s

    def add_analyse(self,s:str):
        self.analyse = self.analyse + s

    def add_come_from(self,s:str):
        self.come = self.come + s

    def format_all(self):
        self.main = self.main.replace("."," .")
        arr = ["A","B","C","D","E","F","G","H"]
        last = False
        for a in arr:
            self.main = self.main.replace( " " + a + ".", "【" + a + "】")

        from_id: int = get_first_number_before_dot(self.answer)
        to_id: int = get_max_number(self.answer)
        answer_arr = []
        for i in range(to_id - from_id + 1):
            tmp_id: int = from_id + i
            r: str = getNumPairs(tmp_id, self.answer, None, to_id)
            answer_arr.append(r.strip()+"=>1")
        self.answer = ";".join(answer_arr)



curState = State.invaid
queId = 0
isValid = False
comprehensionArr = []
curComprehensItem:ComprehensionData = None
curQuestionDataItem:QuestionData = None

def main():
    if len(sys.argv) > 1:
        fileName = sys.argv[1]
    else:
        fileName = getImportPath()
    if not fileName.endswith(".docx"): # 文件夹
        for path in os.listdir(fileName):
            if path.endswith(".docx"):
                path = fileName + "/" + path
                HandleFile(path)
    else:
        HandleFile(fileName)

def formatUnderline(d:docx.text.paragraph.Paragraph):
    for item in d.runs :
        if item.underline :
            item.text = item.text.replace(" ","_")
            item.text = "____________"


def HandleFile(fileName:str):
    global curState,queId,isValid,comprehensionArr,curComprehensItem,curQuestionDataItem
    curState = State.invaid
    queId = 0
    isValid = False
    comprehensionArr = []
    curComprehensItem = None
    curQuestionDataItem = None

    if fileName.__contains__("[上传]"):
        return
    if not os.path.exists(fileName):
        print( "文件不存在%s" % fileName)
        return
    doc = getDoc(fileName)
    print("----------------开始处理文件 %s ------------" % fileName)
    global category
    global type
    name, category, type, mark = AnalyzeFileName(fileName)

    global questionRule
    global optionROle
    # global  curState
    questionRule = re.compile(r"^\d+\s*\.\s*")
    optionROle = re.compile(r"^[A-Z]\s*\.\s*")
    color:any = None
    last_color:any = None
    for d in doc.paragraphs:
        formatUnderline(d)
        str = d.text.replace("．", ".")
        if str == '':
            continue
        tmpValid = isValid
        checkContent(str)
        if tmpValid == False and isValid :
            curState = State.main
            continue
        color = get_paragraph_shading(d)
        if color is None and last_color:
            curState = State.main
            curComprehensItem = None
        last_color = color
        if curState == State.main:
            str = "\n        " + str
        if curState == State.finish :
            break
        if curState == State.invaid:
            continue
        updateState(str)

    ## 写答案
    fileName = fileName.split(".")[0] + "[上传].docx"
    writeAns(fileName)
    print("----------------结束处理文件----------------------------\n")
    return True

def checkContent(s:str):
    global curState
    global isValid
    global curComprehensItem
    global  curQuestionDataItem
    if s.find("//题目开始") != -1:
        isValid = True
        return
    if s.find("//题目结束") != -1:
        isValid = False
        curState = State.finish
        return
    if isValid:
        # 只有题干或选项，才有可能是题目
        if ( curState == State.main or curState == State.opt ) and isQustion(s): # 题目
            curQuestionDataItem = None
            curState = State.que
            return
        if ( curState == State.que or curState == State.opt ) and isOption(s):  # 选项
            curState = State.opt
            return
        if s.find("【答案】") != -1 or s.find("【知识点】") != -1 or s.find("【解析】") != -1 :
            curState = State.answer
            return
        if s.find("【导语】") != -1 or s.find("【解析】") != -1:
            curState = State.analyse
            return
        if s.find("【来源】") != -1:
            curState = State.come_from
            return


def updateState(s:str):
    global curState
    global isValid
    global comprehensionArr
    global curComprehensItem
    global curQuestionDataItem
    if not isValid :
        return;
    match(curState):
        case State.main:
            pattern = r'^#\d+\.'
            m = re.match(pattern, s.strip())
            if curComprehensItem is None:
                curComprehensItem = ComprehensionData()
                comprehensionArr.append(curComprehensItem)
            curComprehensItem.addMainContent(s)
        case State.answer:
            curComprehensItem.add_answer(s)
        case State.come_from:
            curComprehensItem.add_come_from(s)
        case State.analyse:
            curComprehensItem.add_analyse(s)



def isOption(str):
    global optionROle
    if optionROle.match(str):
        return True
    return False

def AnalyzeFileName(str):
    pathArr = str.split("/")
    str = pathArr[len(pathArr)-1]
    titleNameArr = str.split(".")
    if len(titleNameArr) >= 1 :
        str = titleNameArr[0]
    arr = str.split("_")
    if len(arr) == 0 :
        return str,"","",1
    name = arr[0]
    category = ""
    type = ""
    mark = 1
    if len(arr) > 0 :
        category = arr[1]
    if len(arr) > 1 :
        type = arr[2]
    return name,category,type,mark

def isQustion(str):
    if len(str) <= 1 :
        return False
    global questionRule
    if questionRule.match(str):
        return True
    return False

def isAns(str):
    global  analysisRule
    if analysisRule.match(str):
        return True
    return False


def getDoc(filename):
    d = docx.Document(filename)
    return d

def writeAns(fileName):
    '''python-docx实现操作word文档基础命令(包含插入各级标题)
    :param fileName:文件保存路径
    :return: None
    '''

    global category
    global type
    doc = Document()  # 创建新文档
    para1 = doc.add_paragraph()
    content:str = ""
    comNum = 0
    allQue:int = 0
    for c in comprehensionArr:
        comNum += 1
        com:ComprehensionData = c
        com.format_all()
        com.main = com.main + "\n分类:"+type
        content += com.main + "\n"
        content += ("答案:" + com.answer + "\n")
        content += ("解析:" + com.come + "\n")
        content += com.analyse + "\n\n\n\n"



    run_2 = para1.add_run(content)  # 以add_run的方式追加内容，方便后续格式调整
    run_2.font.name = 'Times New Roman'  # 注：这个好像设置 run 中的西文字体
    # 设置中文字体
    # 需导入 qn 模块
    from docx.oxml.ns import qn
    # run_2.font.name = '楷体'  # 注：如果想要设置中文字体，需在前面加上这一句
    run_2.font.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
    # 设置字体大小
    run_2.font.size = Pt(14)
    #run.font.size = Pt(10.5)
    if os.path.exists(fileName):
        os.remove(fileName)
    doc.save(fileName)  # 文档保存
    print("保存完成,题干数量%d,子题目数量%d" % (len(comprehensionArr), allQue) )



if __name__ == '__main__':
    main()
