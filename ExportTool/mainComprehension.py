# coding=UTF-8
import os
import sys
import docx
import re
from docx import Document
# import xlrd
# import xlwt
# from xlutils.copy import copy
import xlwt as xlwt
import xml.etree.ElementTree as ET
from enum import Enum

from docx.opc.oxml import qn
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph

from Tools import getNumPairs, getImportPath, split_options, get_paragraph_shading, iter_block_items, read_table, \
    add_underlines_runs_by_arr


class State(Enum):
    invaid = 0 # 无效
    main = 1 # 题干
    que = 2 # 题目
    opt = 3 # 选项
    answer = 4 # 答案\分类\分数\解析
    finish = 5

class QuestionData:
    id = 0
    oldId = 0
    main = ""
    qus = "题目"
    opt = []
    ans = ""
    analysis = "解析"
    category = "知识点"
    qusType = "题型"
    mark = 1
    isSignle:bool = False

    def __init__(self, ques:str):

        self.id = int(ques.split(".")[0])
        self.oldId = self.id
        self.qus = ques
        self.opt = []
        self.ans = "答案"
        self.analysis = ""
        self.isSignle = False
        self.main = ""
        self._optStr = ""

    def AddMain(self,mainContent):
        self.main = self.main + mainContent

    def AddQueContent(self, ques):
        self.qus = self.qus + ques

    def AddOption(self, opts):
        self._optStr = self._optStr + "\t" + opts
        return True

    def UpdateOpt(self):
        self.opt = split_options(self._optStr)
        if len(self.opt) > 6:
            print("id:%d 选项格式有错(或者下一题的题目格式有错,题目一定是数字+.的格式)，请检查,选项超过4个 %s" % (self.oldId, " ".join(self.opt)))
            return False
        self.opt.sort()
        return True

    def AddAnswer(self, answer):
        if answer.find("【答案】") != -1:
            answer = answer.replace("．",".")
            answer = answer.replace("【答案】","")
            answer = answer.strip()
            # 使用正则表达式匹配数字和对应的答案
            if answer.find(".") != -1:
                matches = re.findall(r'(\d+)\.\s*(\S+)', answer)
                answers = {}
                if matches:
                    for match in matches:
                        number, a = match
                        answers[number] = a

                if not answers.__contains__(str(self.id)):
                    print("id:%d 答案找不到,当前答案是:%s" % (self.id,answer))
                    return
                self.ans = answers[str(self.id)]
            else:
                self.ans = answer
        else:
            self.analysis = self.analysis + answer

    def updateAnl(self,max:int):
        paragraphs = []
        current_paragraph = []
        self.qus = self.qus.replace("("," [").replace("（"," [").replace(")","] ").replace("）","] ")
        self.analysis = self.analysis.replace("【分析】","【解析】").replace("【详解】","【解析】").replace("【导语】","【解析】")
        pattern = r'【解析】(.*?)(?=【)'
        if not self.isSignle and not self.analysis.__contains__("【点睛】"):
            self.analysis = self.analysis + "【"
        matches = re.findall(pattern, self.analysis, re.DOTALL)
        anlDic = {}
        if not self.isSignle and len(matches) == 0:
            print("id=%d 解析中答案查找失败" % self.id)
        for idx, item in enumerate(matches, start=1):
            text = item.strip()
            tmpAnl = getNumPairs(self.id,text,"【",max )
            if not tmpAnl:
                if self.isSignle :
                    return
                print("找不到解析: id = %d" % self.id )
                return
            tmpAnl = re.sub(pattern,"【解析】"+tmpAnl,self.analysis,flags=re.DOTALL)
            self.analysis = tmpAnl
            if self.analysis.endswith("【"):
                self.analysis = self.analysis[:-1]
            return

    def AddAnalysis(self, analysis):
        self.analysis = analysis

    def AddCategory(self,value):
        self.category = value

    def AddQusType(self,qusType):
        self.qusType = qusType

    def AddMark(self,value):
        self.mark = value

    def tostring(self):
        return str(self.id) + "." + self.qus + " 选项:" + " ".join(self.opt) + " 答案:" + self.ans;

class ComprehensionData:
    id:int = 0
    ques:[] = []
    main:str = "" # 题干

    def __init__(self):
        global queId
        queId = queId + 1
        self.id = queId
        self.ques = []
        self.main = ""


    def addMainContent(self,mainContent:str):
        mainContent = mainContent.replace("、",",")
        m = re.match( r'^#\d+\.',mainContent)
        if m :
            print("题干的内容不能以数字+点开头 :" + mainContent)
        self.main = self.main + mainContent

    def addQue(self,que:QuestionData):
        self.ques.append(que)
        tmpId = len(self.ques)
        que.id = tmpId


curState = State.invaid
queId = 0
isValid = False
comprehensionArr = []
curComprehensItem:ComprehensionData = None
curQuestionDataItem:QuestionData = None

def main():
    if len(sys.argv) > 1:
        fileName = sys.argv[1]
    else:
        fileName = getImportPath()
    if not fileName.endswith(".docx"): # 文件夹
        for path in os.listdir(fileName):
            if path.endswith(".docx"):
                path = fileName + "/" + path
                HandleFile(path)
    else:
        HandleFile(fileName)

def formatUnderline(d:docx.text.paragraph.Paragraph,arr:[]):
    for item in d.runs :
        if item.underline :
            s:str = item.text.replace("．", ".")
            item.text = "^" + s + "^"
            arr.append(item.text)



def HandleFile(fileName:str):
    global curState,queId,isValid,comprehensionArr,curComprehensItem,curQuestionDataItem
    curState = State.invaid
    queId = 0
    isValid = False
    comprehensionArr = []
    curComprehensItem = None
    curQuestionDataItem = None

    if fileName.__contains__("[上传]"):
        return
    if not os.path.exists(fileName):
        print( "文件不存在%s" % fileName)
        return
    doc = getDoc(fileName)
    print("----------------开始处理文件 %s ------------" % fileName)
    global category
    global type
    name, category, type, mark = AnalyzeFileName(fileName)

    global questionRule
    global optionROle
    # global  curState
    questionRule = re.compile(r"^\d+\s*\.\s*")
    optionROle = re.compile(r"^[A-Z]\s*\.\s*")
    color:any = None
    last_color:any = None
    under_line_arr = []
    # -----
    for block in iter_block_items(doc):
        str: str = ""
        # 是否为固定选项部分
        is_fix_opt: bool = False
        is_paragraph = False
        if isinstance(block, Paragraph):
            formatUnderline(block,under_line_arr)
            str = block.text
            is_paragraph = True
        elif isinstance(block, docx.table.Table):
            is_fix_opt = True
            str = read_table(block)
        if str == '':
            continue
        tmpValid = isValid
        str = str.replace("．",".")
        checkContent(str)
        if tmpValid == False and isValid :
            curState = State.main
            continue
        color = get_paragraph_shading(block)
        if color is None and last_color:
            curState = State.main
            curComprehensItem = None
        last_color = color
        if curState == State.main:
            str = "\n        " + str
        if curState == State.finish :
            break
        if curState == State.invaid:
            continue
        updateState(str)

    ## 写答案
    fileName = fileName.split(".")[0] + "[上传].docx"
    writeAns(fileName,under_line_arr)
    print("----------------结束处理文件----------------------------\n")
    return True

def checkContent(s:str):
    global curState
    global isValid
    global curComprehensItem
    global  curQuestionDataItem
    if s.find("//题目开始") != -1:
        isValid = True
        return
    if s.find("//题目结束") != -1:
        isValid = False
        curState = State.finish
        return
    if isValid and s.strip().startswith("#"):
        s = s.strip()
        curComprehensItem = None
        curQuestionDataItem = None
        curState = State.main
        return
    if isValid:
        # 只有题干或选项，才有可能是题目
        if ( curState == State.main or curState == State.opt ) and isQustion(s): # 题目
            curQuestionDataItem = None
            curState = State.que
            return
        if ( curState == State.que or curState == State.opt ) and isOption(s):  # 选项
            curState = State.opt
            return
        if s.find("【答案】") != -1 or s.find("【知识点】") != -1 or s.find("【解析】") != -1 :
            curState = State.answer
            return

def updateState(s:str):
    global curState
    global isValid
    global comprehensionArr
    global curComprehensItem
    global curQuestionDataItem
    if not isValid :
        return;
    match(curState):
        case State.main:
            pattern = r'^#\d+\.'
            m = re.match(pattern, s.strip())
            if curComprehensItem is None:
                curComprehensItem = ComprehensionData()
                comprehensionArr.append(curComprehensItem)
            if s.find("#") != -1:
                if m:
                    mlen = len(m.group(0))
                    ss = s[mlen:]
                    curComprehensItem.addMainContent(ss)
                else:
                    curComprehensItem.addMainContent(s.split("#")[-1])
            else:
                curComprehensItem.addMainContent(s)
            # 如果题干和题目是一起
            if m:
                curQuestionDataItem = QuestionData(s.split("#")[-1])
                curQuestionDataItem.isSignle = True
                curComprehensItem.ques.append(curQuestionDataItem)
        case State.que:
            s = s.replace("、", ",")
            if curQuestionDataItem is None:
                curQuestionDataItem = QuestionData(s)
                curComprehensItem.ques.append(curQuestionDataItem)
            else:
                curQuestionDataItem.AddQueContent(s)
        case State.opt:
            if curQuestionDataItem is None:
                print("没有题目，直接就是选项目，请检查题干为:%s " % curComprehensItem.main)
                return
            curQuestionDataItem.AddOption(s)
        case State.answer:
            for item in curComprehensItem.ques:
                item.AddAnswer(s)



def isOption(str):
    global optionROle
    if optionROle.match(str):
        return True
    return False

def AnalyzeFileName(str):
    pathArr = str.split("/")
    str = pathArr[len(pathArr)-1]
    titleNameArr = str.split(".")
    if len(titleNameArr) >= 1 :
        str = titleNameArr[0]
    arr = str.split("_")
    if len(arr) == 0 :
        return str,"","",1
    name = arr[0]
    category = ""
    type = ""
    mark = 1
    if len(arr) > 0 :
        category = arr[1]
    if len(arr) > 1 :
        type = arr[2]
    return name,category,type,mark

def isQustion(str):
    if len(str) <= 1 :
        return False
    global questionRule
    if questionRule.match(str):
        return True
    return False

def isAns(str):
    global  analysisRule
    if analysisRule.match(str):
        return True
    return False


def getDoc(filename):
    d = docx.Document(filename)
    return d

def writeAns(fileName,under_line_arr:[]):
    '''python-docx实现操作word文档基础命令(包含插入各级标题)
    :param fileName:文件保存路径
    :return: None
    '''

    global category
    global type
    doc = Document()  # 创建新文档
    para1 = doc.add_paragraph()
    content:str = ""
    comNum = 0
    allQue:int = 0
    for c in comprehensionArr:
        comNum += 1
        com:ComprehensionData = c
        if len(com.ques) == 0 :
            print("题目数量为0:题干:" + com.main)
        allQue += len(com.ques)
        if com.main == "" :
            if len(com.ques) > 1:
                print("无题干:题目" + com.ques[0].qus)
            else:
                com.main = com.ques[0].qus
        com.main = com.main.replace("("," [").replace("（"," [").replace(")","] ").replace("）","] ")
        if bool(re.match(r'^\d+\.', com.main.strip())):
            print("题干中，不能以 \"数字+.\" 开头，如\"1.\" ，请将\".\"改成其他字符 , 题干 =  %s " % com.main )
        content += "[理解题开始]\n"
        com.main = com.main.strip() + "\n分类:"+type
        content += ( com.main + " (" + str(len(com.ques)) + "题)\n\n\n" )
        num = 0
        for q in com.ques:
            que:QuestionData = q
            que.UpdateOpt()
            que.updateAnl( com.ques[0].id + len(com.ques) - 1)
            if que.ans == "":
                print("答案为空,id=%d" % que.id )
            # 题目
            content += que.qus + "\n"
            for op in que.opt:
                if len(op.strip()) <= 2:
                    print("选项为空，请注意检查 id = %d,选项:%s" % (que.id,op) )
                content += op + "\n"
            content += "答案:" + que.ans + "\n"
            content += "分数:5\n"
            que.analysis = que.analysis.replace("【","\n【")
            num += 1
            if num == len(com.ques):
                content += "解析:" + que.analysis + "\n\n"
            else:
                content += "解析:" + que.analysis + "\n\n\n"
        if comNum == len(comprehensionArr):
            content += "[理解题结束]"
        else:
            content += "[理解题结束]\n"

    if len(under_line_arr) > 0 :
        add_underlines_runs_by_arr(para1,content)
    else:
        run_2 = para1.add_run(content)  # 以add_run的方式追加内容，方便后续格式调整
        run_2.font.name = 'Times New Roman'  # 注：这个好像设置 run 中的西文字体
        # 设置中文字体
        # 需导入 qn 模块
        from docx.oxml.ns import qn
        # run_2.font.name = '楷体'  # 注：如果想要设置中文字体，需在前面加上这一句
        run_2.font.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
        # 设置字体大小
        run_2.font.size = Pt(14)
    if os.path.exists(fileName):
        os.remove(fileName)
    doc.save(fileName)  # 文档保存
    print("保存完成,题干数量%d,子题目数量%d" % (len(comprehensionArr), allQue) )



if __name__ == '__main__':
    main()
