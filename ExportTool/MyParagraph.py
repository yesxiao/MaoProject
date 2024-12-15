import os
import shutil
import zipfile
import xml.etree.ElementTree as ET
from docx import Document, ImagePart
from docx.oxml import CT_Drawing
from docx.shared import Pt
from docx.text.run import Run

from Tools import get_paragraph_shading


#获得多个以start,end起始的字符串
def find_all(content:str,start:str,end:str):
    data_list = []
    find_idx = 0
    while True:
        index_start = content.find(start,find_idx)
        index_end = -1
        if index_start != -1 :
            index_end = content.find(end,index_start+len(start))
            if index_end != -1:
                c:str = content[index_start+len(start):index_end]
                c_index:int = index_start
                o = ContentItem(c_index,c)
                data_list.append(o)
                find_idx = index_end + 1
            continue
        break
    return data_list

class ContentItem:
    def __init__(self,index,content):
        self.index = index
        self.content = content
        self.img_path = None
        self.img_size_x = 0
        self.img_size_y = 0

    #删除图片相关描述文本
    def del_img_describle(self,id2path:any):
        path = id2path[self.content]
        if path :
            self.img_path = path
            self.content = ( "[path:\"%s\"]" % path )



class ParagraphsAnlyse:

    def __init__(self,paragraph:any,id2path:any):
        self.paragraph = paragraph ;
        paragraphXml: str = paragraph._element.xml
        arr_text = find_all(paragraphXml, "<w:t>", "</w:t>")
        self.contents = arr_text
        if self.contents is None:
            self.contents = []
        #arr_drawing
        arr_drawing = find_all(paragraphXml, "<w:drawing>", "</w:drawing>")
        for drawing in arr_drawing:
            arr_size = find_all(drawing.content,"<wp:extent cx=\"","\"/>")
            size_x:int = 0
            size_y:int = 0
            if len(arr_size) > 0 :
                size_str:str = arr_size[0].content
                size_str_arr = size_str.split("\" cy=\"")
                size_x = int(size_str_arr[0])
                size_y = int(size_str_arr[1])
            arr_img = find_all(drawing.content, "<a:blip r:embed=\"", "\"/>")
            for img in arr_img:
                img.del_img_describle(id2path)
                img.index = drawing.index
                img.img_size_x = size_x
                img.img_size_y = size_y
                self.contents.append(img)
        self.contents.sort(key=lambda c: c.index)
        self.color = get_paragraph_shading(paragraph)

class DocAnalyse:

    def __init__(self,doc_path:str):
        self.paragraph_arr, self.unpack_dir ,self.doc = read_doc(doc_path)

def unpack_docx_file(file_name:str):
    zip_path: str = file_name.replace(".docx",".zip")
    unpack_dir:str = file_name.replace(".docx","")

    if os.path.exists(zip_path):
        os.remove(zip_path)
    if os.path.exists(unpack_dir):
        shutil.rmtree(unpack_dir)
    os.mkdir(unpack_dir)
    shutil.copyfile( file_name, zip_path)
    f = zipfile.ZipFile(zip_path, 'r')
    # 将图片提取并保存
    for file in f.namelist():
        f.extract(file, unpack_dir)
    # 释放该zip文件
    f.close()
    os.remove(zip_path)
    return unpack_dir

def read_img_relation(unpack_dir,doc):
    xml_path:str = unpack_dir+"/word/_rels/document.xml.rels"
    tree = ET.parse(xml_path)
    root = tree.getroot()
    dic = {}
    for element in root:
        id = element.get("Id")
        path = element.get("Target")
        dic[id] = path
        part = doc.part.related_parts[id]
        if isinstance(part, ImagePart):
            img:ImagePart = part

    return dic


def read_doc(doc_path:str):
    doc = Document(doc_path)
    unpack_dir = unpack_docx_file(doc_path)
    id2path = read_img_relation(unpack_dir,doc)
    plist = []
    for para in doc.paragraphs:
        p = ParagraphsAnlyse(para, id2path)
        plist.append(p)
    return plist,unpack_dir,doc

#test
if __name__ == '__main__':
    doc_path = r"D:\猫猫\0905\肖鲲测试\【高中地理】03地球上的大气（06气压带）风带与气候）（02）101~200_选择题_1_1.docx"
    save_doc_path = r"D:\猫猫\0905\肖鲲测试\【高中地理】03地球上的大气（06气压带）风带与气候）（02）101~200_选择题_1_1[save].docx"
    doc_analyse = DocAnalyse(doc_path)
    paragraphs_arr = doc_analyse.paragraph_arr
    unpack_dir = doc_analyse.unpack_dir
    doc = doc_analyse.doc
    from docx.oxml.ns import qn
    for p in paragraphs_arr:
        pa = doc.add_paragraph()
        for c in p.contents:
            run = pa.add_run()
            run.font.name = 'Times New Roman'
            # run_2.font.name = '楷体'  # 注：如果想要设置中文字体，需在前面加上这一句
            run.font.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
            # 设置字体大小
            run.font.size = Pt(14)
            tmp_content:str = c.content
            run.text = tmp_content
            if run.text.startswith("[path:\""):
                path = tmp_content.replace("[path:\"","").replace("\"]","")
                run.text = ""
                r:Run = run
                r.add_picture(unpack_dir+ "/word/" +path,c.img_size_x,c.img_size_y)
    if os.path.exists(unpack_dir):
        shutil.rmtree(unpack_dir)
    doc.save(save_doc_path)  # 文档保存